import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Mirar para intervenir: Una reflexión situada sobre "Gloriosa Victoria"', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction in Steiman tone
    p1 = doc.add_paragraph()
    p1.add_run('Me detengo frente a la imagen. No es solo un archivo en el escritorio; es una ventana que se abre a una de nuestras tantas heridas latinoamericanas. Como dice Jorge Steiman (aunque no lo cito, su voz resuena en mi cabeza), la enseñanza es una práctica compleja, y mirar el arte de Diego Rivera es, en sí mismo, un acto pedagógico que nos obliga a tomar posición. No se puede ser neutral frente a este lienzo.').italic = True

    # Body
    doc.add_heading('La escena del crimen', level=1)
    
    p2 = doc.add_paragraph()
    p2.add_run('En "Gloriosa Victoria" (1954), Rivera no nos propone una contemplación pasiva. Nos arroja al centro de la ignominia. Veo ese apretón de manos central: John Foster Dulles y el militar Castillo Armas. Es el rito del pacto, la formalización del "contrato didáctico" de la dominación, donde lo que se enseña es la sumisión.')
    
    p3 = doc.add_paragraph()
    p3.add_run('La mirada se desvía inevitablemente hacia la bomba. Esa cara de Dwight D. Eisenhower grabada en el metal frío. Es la metáfora de una victoria que se construye sobre la destrucción de la democracia guatemalteca. Alrededor, los cajones de bananas de la United Fruit Company. No son solo frutas; son los símbolos del poder económico que dicta la historia desde las sombras.')
    
    doc.add_heading('El barro de la historia', level=1)
    
    p4 = doc.add_paragraph()
    p4.add_run('Abajo, en el primer plano, los cuerpos. Los trabajadores, los niños, el pueblo que pone el cuero mientras los de arriba se saludan. Rivera nos muestra que la historia no sale de los escritorios pulcros, sino que se amasa en el barro, en el dolor y en la resistencia.')
    
    p5 = doc.add_paragraph()
    p5.add_run('Esta imagen nos interpela hoy. ¿Qué estamos mirando cuando miramos la historia? ¿Qué contratos estamos aceptando o rechazando en nuestras propias prácticas? Mirar este mural con una "mirada situada" es reconocer que somos parte de este continente, que estas contradicciones nos habitan y que, como docentes y estudiantes, nuestra "Gloriosa Victoria" no debería ser la del opresor, sino la de la memoria y la verdad.')

    # Closing
    p_final = doc.add_paragraph()
    p_final.add_run('Porque, al final del día, lo que queda no es el apretón de manos de los poderosos, sino la pincelada valiente que se atrevió a decir: esto también pasó.').bold = True

    # Save path
    desktop_path = os.path.join(os.environ['USERPROFILE'], 'Desktop')
    file_path = os.path.join(desktop_path, 'Reflexion_Mural_Gloriosa_Victoria.docx')
    
    doc.save(file_path)
    print(f"Documento creado en: {file_path}")

if __name__ == "__main__":
    create_document()
