import docx

doc = docx.Document()

doc.add_heading('Las invasiones inglesas de 1806 y 1807', 0)
doc.add_paragraph('Jorge Medina - 08/04/2026 23:33')

doc.add_paragraph('¡Hola a todos y todas! Comparto mi aporte para el foro de esta primera clase. Más que repasar los hechos de manera puramente cronológica, me parece interesante tratar de desarmarlos y pensar cómo condicionaron el escenario político posterior, a partir de la lectura de Marcela Ternavasio.')

p = doc.add_paragraph('Respondiendo al interrogante, creo que podemos pensar el desarrollo en tres grandes ejes que se interconectan directamente:')

doc.add_heading('1. El quiebre del proyecto defensivo borbónico', level=2)
doc.add_paragraph('Como bien marca la autora, las Invasiones Inglesas significaron, en los hechos, el colapso del plan reformista borbónico en el Río de la Plata. La creación del Virreinato en 1776 tuvo un propósito centralmente estratégico y defensivo: proteger ese rincón del imperio del avance extranjero. Sin embargo, en el momento de la verdad (el desembarco británico en 1806), todo el andamiaje burocrático y militar demostró ser ineficaz. La defensa falló de forma estrepitosa debido a la falta de tropas regulares y a decisiones tácticas erráticas.')

doc.add_heading('2. La grieta vertical y la crisis de Sobremonte', level=2)
doc.add_paragraph('La reacción del virrey Sobremonte de abandonar Buenos Aires hacia Córdoba (llevándose el tesoro y dejando la resistencia a cargo de una Audiencia y un Cabildo que capitularon rápido) fue un golpe durísimo a la figura virreinal. Fue leído por la población como un acto de cobardía y abandono. Ternavasio describe esto como una "grieta vertical en el orden colonial" que hirió de muerte a la máxima autoridad. El hecho sin precedentes de que, a través de un Cabildo Abierto, los vecinos lograran suspenderlo de su mando militar para delegarlo en Liniers, rompió el principio de autoridad intocable que venía de la metrópoli.')

doc.add_heading('3. Las milicias urbanas como nuevo centro del poder', level=2)
doc.add_paragraph('Ante el notorio vacío que dejaron las autoridades peninsulares —incapaces de defender su propio territorio o dispuestas a pactar con el invasor—, la población local tuvo que organizarse desde cero. La militarización urgente dio origen a las milicias urbanas, integradas por vecinos, criollos y españoles, pero sostenidas y financiadas localmente. Este es el verdadero desplazamiento del poder: al lograr reconquistar y defender su ciudad por mérito propio, estas milicias y sus jefes pasaron a tener un peso político decisivo. El apoyo y la fuerza armada ya no dependían de la corona y su burocracia, sino de contingentes armados de base local que, llegado el año 1810, tendrían la última palabra para respaldar la formación de un gobierno propio.')

doc.add_paragraph('Es fascinante ver cómo un conflicto externo y militar termina destapando semejante crisis institucional de fondo y da pie a una nueva configuración del poder tan distinta.')
doc.add_paragraph('¿Qué piensan ustedes, compañeros/as? ¿Pudieron identificar también cómo las milicias fueron ese primer "ensayo" real de participación y decisión política autónoma?')
doc.add_paragraph('¡Saludos para todos!')

doc.save(r'C:\Users\alanm\Desktop\Jorge\ARG 1\Las invasiones inglesas de 1806 y 1807.docx')
