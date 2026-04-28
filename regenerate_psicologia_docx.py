from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pathlib import Path


def u(text):
    return text.encode("ascii").decode("unicode_escape")


out = Path.home() / "Desktop" / "Psicologia_y_Educacion_respuestas_corregido.docx"

doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(u("Psicolog\\u00eda y Educaci\\u00f3n"))
r.bold = True
r.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(u("S\\u00edntesis con referencias de p\\u00e1gina"))
r.italic = True

intro = doc.add_paragraph()
intro.add_run(u("Fuente: ")).bold = True
intro.add_run(
    u(
        "Psicolog\\u00eda y educaci\\u00f3n: una relaci\\u00f3n indiscutible "
        "(Susana Leliwa e Irene Scangarello). Las p\\u00e1ginas indicadas "
        "corresponden al OCR del PDF escaneado del escritorio."
    )
)

items = [
    (
        u("\\u00bfQu\\u00e9 es la Psicolog\\u00eda?"),
        u(
            "Es la ciencia que estudia la conducta y los fen\\u00f3menos "
            "internos de la experiencia humana. Tambi\\u00e9n analiza lo que "
            "las personas hacen, sienten y piensan seg\\u00fan su modo de ser y "
            "sus circunstancias. No se ocupa solo de lo observable, sino "
            "tambi\\u00e9n de lo subjetivo y mental."
        ),
        "pp. 20-23.",
    ),
    (
        u("\\u00bfQu\\u00e9 aporte brinda la Psicolog\\u00eda a la educaci\\u00f3n?"),
        u(
            "Ayuda a comprender c\\u00f3mo aprende y se desarrolla el sujeto en "
            "situaciones educativas. Aporta claves sobre motivaci\\u00f3n, "
            "madurez, conocimientos previos, v\\u00ednculos y contexto social. "
            "Sirve para planificar mejor la ense\\u00f1anza y mejorar la pr\\u00e1ctica "
            "docente."
        ),
        "pp. 21-23.",
    ),
    (
        u(
            "\\u00bfC\\u00f3mo llega a constituirse la Psicolog\\u00eda como ciencia independiente? "
            "\\u00bfQu\\u00e9 problem\\u00e1ticas tuvo?"
        ),
        u(
            "Se independiza de la Filosof\\u00eda a fines del siglo XIX, con el "
            "laboratorio de Wundt en 1879. Su problema central fue definir su "
            "objeto de estudio: alma, conciencia, conducta o inconsciente. "
            "Tambi\\u00e9n debati\\u00f3 la tensi\\u00f3n entre objetividad y subjetividad "
            "al estudiar al ser humano."
        ),
        "pp. 24-25.",
    ),
    (
        u(
            "\\u00bfQu\\u00e9 son los campos de la Psicolog\\u00eda? "
            "\\u00bfQu\\u00e9 nos permite darnos cuenta eso?"
        ),
        u(
            "Son las distintas \\u00e1reas en que la Psicolog\\u00eda investiga, como "
            "cl\\u00ednica, educativa o social. Muestran que la disciplina se "
            "especializ\\u00f3 mucho y se volvi\\u00f3 m\\u00e1s diversa y compleja. "
            "Tambi\\u00e9n dejan ver que la Psicolog\\u00eda no es una sola mirada, "
            "sino varias en relaci\\u00f3n."
        ),
        "p. 20.",
    ),
    (
        u(
            "\\u00bfEl accionar docente se basa en experiencia, vocaci\\u00f3n o fundamentos psicol\\u00f3gicos?"
        ),
        u(
            "El libro sostiene que no alcanza solo con experiencia o vocaci\\u00f3n. "
            "El docente necesita fundamentos psicol\\u00f3gicos para tomar mejores "
            "decisiones sobre ense\\u00f1anza y evaluaci\\u00f3n. Eso le permite "
            "comprender al alumno y revisar cr\\u00edticamente su propia pr\\u00e1ctica."
        ),
        "pp. 22-23.",
    ),
    (
        u("\\u00bfQu\\u00e9 aporte al campo educativo han brindado los psic\\u00f3logos?"),
        u(
            "Brindaron teor\\u00edas para entender el aprendizaje, el desarrollo, la "
            "motivaci\\u00f3n y la inteligencia. Tambi\\u00e9n ayudaron a pensar el "
            "v\\u00ednculo docente-alumno y los factores que influyen en la "
            "ense\\u00f1anza. Sus aportes permiten mejorar la educaci\\u00f3n desde "
            "una base cient\\u00edfica."
        ),
        "pp. 24-25 y 30-33.",
    ),
    (
        u(
            "\\u00bfCu\\u00e1l es la relaci\\u00f3n entre Psicolog\\u00eda del desarrollo y Psicolog\\u00eda educacional?"
        ),
        u(
            "Las dos estudian cambios en el comportamiento humano, pero desde "
            "enfoques distintos. La Psicolog\\u00eda del desarrollo mira los cambios "
            "como parte del crecimiento humano. La Psicolog\\u00eda educacional los "
            "entiende como resultado de la ense\\u00f1anza y el aprendizaje."
        ),
        "p. 22.",
    ),
    (
        u("\\u00bfA qu\\u00e9 hacemos referencia cuando hablamos de complejidad del acto educativo?"),
        u(
            "A que educar no depende de un solo factor, sino de muchos al mismo "
            "tiempo. Intervienen lo biol\\u00f3gico, lo cognitivo, lo afectivo, lo "
            "social, lo cultural y lo institucional. Por eso ense\\u00f1ar y "
            "aprender son procesos complejos y no lineales."
        ),
        "pp. 21-23.",
    ),
    (
        u(
            "Cuando hablamos de desarrollo humano, \\u00bfqu\\u00e9 factores debemos tener en cuenta?"
        ),
        u(
            "Debemos considerar procesos madurativos, pero tambi\\u00e9n el contexto "
            "social e hist\\u00f3rico. El desarrollo humano ocurre en relaci\\u00f3n "
            "con la cultura, la escuela, la familia y los otros. No es solo "
            "biol\\u00f3gico: se construye permanentemente en interacci\\u00f3n con el entorno."
        ),
        "p. 22.",
    ),
]

for i, (question, answer, ref) in enumerate(items, 1):
    pq = doc.add_paragraph()
    rq = pq.add_run(f"{i}. {question}")
    rq.bold = True
    doc.add_paragraph(answer)
    pr = doc.add_paragraph()
    pr.add_run(u("Referencia: ")).bold = True
    pr.add_run(ref)

doc.save(str(out))
print(out)
